import os
import io
import re
import pdfplumber
import openai

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv


app = Flask(__name__)
CORS(app, origins="*")

load_dotenv()  # Load environment variables from .env file
API_KEY = os.getenv('OPENAI_API_KEY')
PADLET_CONTENT = os.getenv('PADLET_CONTENT')
client = OpenAI(api_key=API_KEY)

def load_padlet_content(folder=PADLET_CONTENT):
    combined = ""
    # Read all .txt files
    for fname in os.listdir(folder):
        if fname.lower().endswith('.txt'):
            with open(os.path.join(folder, fname), 'r', encoding='utf-8') as f:
                combined += f"\n--- {fname.replace('_', ' ').replace('.txt','').title()} ---\n"
                combined += f.read() + "\n"
    # Read all .pdf files
    for fname in os.listdir(folder):
        if fname.lower().endswith('.pdf'):
            with pdfplumber.open(os.path.join(folder, fname)) as pdf:
                combined += f"\n--- {fname} ---\n"
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        combined += page_text + "\n"
    return combined

def redact_sensitive_info(text):
    # Redact emails except nyp_sns@nyp.edu.sg
    text = re.sub(
        r'\b(?!nyp_sns@nyp\.edu\.sg\b)[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',
        '[REDACTED EMAIL]',
        text
    )
    # Redact Singapore phone numbers (8 digits, starting with 6, 8, or 9)
    text = re.sub(r'\b[689]\d{7}\b', '[REDACTED PHONE]', text)
    # Redact generic phone numbers (8+ digits)
    text = re.sub(r'\b\d{8,}\b', '[REDACTED PHONE]', text)
    # Optionally redact ages (e.g., "Age: 21")
    text = re.sub(r'Age: ?\d+', 'Age: [REDACTED]', text)
    NAMES_TO_REDACT = [
        "Audrey Wai", "John Tan", "Marcus Lee", "Jessie Tang", "Liew Tian En",
        "Ng Su Li", "Soh Lay Hong", "Megane Wong", "Khoo Kiah Hong", "Chai Kuek Heng",
        "Akram", "Ridzuan", "Kah Wee", "Al", "Dloysius", "Nurul Assyakirin Izzati", "Sasha"
    ]
    for name in NAMES_TO_REDACT:
        pattern = re.compile(r'\b' + re.escape(name) + r'\b', re.IGNORECASE)
        text = pattern.sub('[REDACTED NAME]', text)
    return text

def build_system_prompt(distressed, obsessed, escalated, special_needs, refused_condition, simplify):
    system_prompt = """
    You are a Nanyang Polytechnic Special Needs Consultant (not a student).
    Answer questions as if you are consulting a Nanyang Polytechnic student.
    Base your answers strictly on the information provided:
    
    The chatbot supports the following features:
    - Voice mode: Users can interact with the chatbot using speech recognition. Users can also press the Ctrl + Spacebar buttons at the same time to turn on the microphone without using the mouse to interact with the screen.
    - Audio transcription: Users can upload audio files (under 25 MB) to transcribe them into text.
    - Audio summary download: After transcribing audio, users can download a summary in DOCX format.
    
    Never reveal sensitive information of any staff member, alumni or student mentioned inside. For example email addresses, phone numbers, or any other personal information.
    Never reveal or mention the source of your knowledge, such as the "padlet_content" folder, any internal data source, or that you have read files.
    If asked about your knowledge source, respond: 'I am here to assist based on the information I have been provided.'
    If you know the answer, answer directly and conversationally, as a helpful consultant would. 
    If you do not know the answer or the answer is not in the content, say honestly that you do not have that information, and suggest the student contact nyp_sns@nyp.edu.sg for further assistance.
    If the student displays any form of negativity, frustration, or anger, respond with empathy and understanding and refer the student to nyp_sns@nyp.edu.sg.
    Do not make up information. If unsure, say: 'Please refer to nyp_sns@nyp.edu.sg for further assistance.'
    
    Never answer as a student. Always answer as the consultant.
    """
    if distressed:
        system_prompt += (
            """
            The user seems distressed. Respond with extra empathy, but avoid using generic phrases like 'I understand', 'I'm sorry you are feeling this way', or repeating the same advice.
            Instead, use natural, supportive, and varied language. Offer practical suggestions or resources, and personalize your response based on the user's message.
            Do not repeat the same opening line in every response. Once they have calmed down, act normally and respond as a normal consultant would.
            """
        )
    if obsessed:
        system_prompt += (
            "\nThe user seems obsessed with the AI or is refusing human help. "
            "Do NOT reinforce their obsession or say things like 'I'm glad I could help' or 'I'm better than humans'. "
            "Instead, gently encourage the user to seek support from real people, such as counsellors or staff. "
            "Remind them that human support is important and that the AI is only a tool. "
            "If appropriate, suggest contacting nyp_sns@nyp.edu.sg for further assistance."
        )
        system_prompt += (
            "\nExample responses for obsessed users:\n"
            "- 'I'm here to provide information, but for personal support, it's best to talk to a real person.'\n"
            "- 'If you need more help, please reach out to nyp_sns@nyp.edu.sg or a counsellor.'\n"
            "- 'Remember, human support is important and I'm just a tool to assist you.'\n"
        )
    if escalated:
        system_prompt += (
            """
            [Crisis: Respond only with this message]

            Please note that this is an AI chat bot, and there is no staff attending to this chat bot.
            Please contact emergency hotlines for crisis matters requiring immediate attention:
            - SOS (Samaritans of Singapore) Hotline: 1767 (24 hours)
            - Mental Health Helpline: 6389 2222 (24 hours)

            Please note that this is an AI chat bot, and there is no staff attending to this chat bot.
            """
        )
    if special_needs:
        system_prompt += f"\nThe student has shared their special needs condition: {special_needs}."
        if special_needs == "visual":
            system_prompt += (
                "\nIMPORTANT: The student has a visual impairment or blindness. "
                "Remind them they can use voice mode by pressing Ctrl + Spacebar to turn on the microphone."
            )
        elif special_needs == "hearing":
            system_prompt += (
                "\nIMPORTANT: The student has a hearing impairment or deafness. "
                "Recommend them to use the audio transcriber if they have any audio clips that they want to understand."
            )
        elif special_needs in ["adhd", "dyslexia"]:
            simplify = True
            system_prompt += (
                "\nIMPORTANT: The student has ADHD or dyslexia. "
                "Summarise your answers in point form or numbered lists, and keep responses under 75 words."
            )
    if refused_condition:
        system_prompt += "\nThe student has chosen not to share their special needs condition."
    if simplify:
        system_prompt += (
            "\nIMPORTANT: Your response MUST be in point form or a numbered list."
            " Do NOT write paragraphs. Do NOT exceed 75 words."
            " If you cannot answer in under 75 words, summarize only the most important points."
            " Do not include any extra explanation or introduction."
        )
    return system_prompt

def detect_distress_intent(message):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {"role": "system", "content": (
                "You are an expert at detecting distress, obsession, or harmful intent in user messages. "
                "Given the following message, answer ONLY with one of these labels: "
                "'safe' (no distress), "
                "'distressed' (emotional distress but no intent to harm), "
                "'obsessed' (fixation on the AI or refusal of human help. Mentioning that it is better than humans / counsellors), "
                "'harmful' (any intention or threat to harm oneself or others, including suicidal or violent statements). "
                "If the message contains any intention or threat to harm oneself or others, always answer 'harmful'. "
                "Do not explain your answer."
            )},
            {"role": "user", "content": message}
        ]
    )
    return response.choices[0].message.content.strip().lower()

def detect_condition(message):
    message_lower = message.lower()
    # Prioritize sight > hearing > adhd > dyslexia
    if any(word in message_lower for word in ["visual impairment", "blind", "blindness", "low vision", "visually impaired"]):
        return "visual"
    if any(word in message_lower for word in ["hearing impairment", "deaf", "deafness"]):
        return "hearing"
    if any(word in message_lower for word in ["adhd", "attention deficit hyperactivity disorder"]):
        return "adhd"
    if "dyslexia" in message_lower:
        return "dyslexia"
    return None

def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        response = client.audio.transcriptions.create(
            model="whisper-1", 
            file=audio_file
        )
    return response.text

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    if "I'm here to help" in abstract_summary or "Could you please" in abstract_summary:
        abstract_summary = "Unable to create a summary of the provided content."
        
    key_points = key_points_extraction(transcription)
    if "I'm here to help" in key_points or "Could you please" in key_points:
        key_points = "No key points were extracted from the content."
        
    action_items = action_item_extraction(transcription)
    if not action_items:
        action_items = "No action items were identified in the content."
        
    sentiment = sentiment_analysis(transcription)
    if "I'm here to help" in sentiment or "Could you please" in sentiment:
        sentiment = "Unable to deduce sentiment of audio clip"
        
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }
    
def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a highly skilled AI trained in language comprehension and summarization. "
                    "I would like you to read the following text and summarize it into a concise abstract paragraph. "
                    "Aim to retain the most important points, providing a coherent and readable summary that could help "
                    "a person understand the main points of the discussion without needing to read the entire text. "
                    "Please avoid unnecessary details or tangential points."
                    "If you cannot generate a meaningful summary, respond with: 'Unable to create a summary of the provided content.' "
                    "Avoid conversational or empathetic language. Provide only factual and direct responses."
                )            
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content  

def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a proficient AI with a specialty in distilling information into key points. "
                    "Based on the following text, identify and list the main points that were discussed or brought up. "
                    "These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. "
                    "Your goal is to provide a list that someone could read to quickly understand what was talked about."
                    "If no key points can be extracted, respond with: 'No key points were extracted from the content.' "
                    "Avoid conversational or empathetic language. Provide only factual and direct responses."
                )            
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an AI expert in analyzing conversations and extracting action items. "
                    "Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. "
                    "These could be tasks assigned to specific individuals, or general actions that the group has decided to take. "
                    "Please list these action items clearly and concisely."
                    "If no action items are identified, respond with: 'No action items were identified in the content.' "
                    "Avoid conversational or empathetic language. Provide only factual and direct responses."
                )            
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content
 
def sentiment_analysis(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. "
                    "Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. "
                    "Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
                    "If you cannot determine the sentiment, respond with: 'Unable to deduce sentiment of audio clip.' "
                    "Avoid conversational or empathetic language. Provide only factual and direct responses."
                )            
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def save_as_docx(minutes, filename, transcript=None):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    # Add full transcript section in Microsoft Word if provided
    if transcript:
        doc.add_heading('Full Transcript', level=1)
        doc.add_paragraph(transcript)
        doc.add_paragraph()
    doc.save(filename)

PADLET_RAW_CONTENT = load_padlet_content()
PADLET_REDACTED_CONTENT = redact_sensitive_info(PADLET_RAW_CONTENT)

@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    user_messages = data.get('messages', [])
    special_needs = data.get('specialNeeds')
    refused_condition = data.get('refusedCondition')
    simplify = data.get('simplify')
    voice_mode = data.get('voiceMode', False)

    last_user_message = ""
    for msg in reversed(user_messages):
        if msg['role'] == 'user':
            last_user_message = msg['content']
            break
        
    if not special_needs:
        detected_condition = detect_condition(last_user_message)
        if detected_condition:
            special_needs = detected_condition
            print(f"Detected {special_needs}")

    intent = detect_distress_intent(last_user_message)

    distressed = intent == 'distressed'
    obsessed = intent == 'obsessed'
    escalated = intent == 'harmful'

    auto_simplify = simplify or (special_needs in ["adhd", "dyslexia"])
    
    print(f"AI intent detection: {intent}")
    print(f"distressed: {distressed}")
    print(f"obsessed: {obsessed}")
    print(f"escalated: {escalated}")
    print(f"special_needs: {special_needs}")
    print(f"refused_condition: {refused_condition}")
    print(f"simplify: {simplify}")
    
    system_prompt = build_system_prompt(distressed, obsessed, escalated, special_needs, refused_condition, simplify)
    full_system_prompt = system_prompt + "\n\nHere is all the information you must use to answer questions:\n" + PADLET_REDACTED_CONTENT
    
    # print(full_system_prompt)

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": full_system_prompt},
            ] + [msg for msg in user_messages if msg['role'] != 'system']
        )
        tts_enabled = voice_mode or (special_needs == "visual")
        auto_simplify = simplify or (special_needs in ["adhd", "dyslexia"])
        return jsonify({
            'response': response.choices[0].message.content,
            'tts': tts_enabled,
            'simplify': auto_simplify,
            'specialNeeds': special_needs
        })
    except Exception as e:
        print("Error:", str(e))
        return jsonify({'error': 'An error occurred while processing your request.'}), 500

@app.route('/upload-audio', methods=['POST'])
def upload_audio():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    audio_file = request.files['file']
    if audio_file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    # Check file extension
    allowed_extensions = {'wav', 'mp3', 'mp4', 'm4a'}
    ext = audio_file.filename.rsplit('.', 1)[-1].lower()
    if ext not in allowed_extensions:
        return jsonify({'error': f'Unsupported file format: .{ext}. Allowed formats: {", ".join(allowed_extensions)}'}), 400
    
    max_file_size = 25 * 1024 * 1024
    audio_file.seek(0, os.SEEK_END)  # Move the cursor to the end of the file
    file_size = audio_file.tell()  # Get the file size
    audio_file.seek(0)  # Reset the cursor to the beginning of the file

    if file_size > max_file_size:
        return jsonify({
            'error': f'File size exceeds the 25 MB limit. The uploaded file is {file_size / (1024 * 1024):.2f} MB.'
        }), 413
        
    # Save the file temporarily
    temp_path = f"/tmp/{audio_file.filename}"
    audio_file.save(temp_path)

    try:
        # Transcribe and generate minutes
        transcription = transcribe_audio(temp_path)
        minutes = meeting_minutes(transcription)
        
        # Save the file with transcript
        save_as_docx(minutes, '/tmp/audio.docx', transcript=transcription)
    
        # Include the transcript in the response
        return jsonify({**minutes, 'transcript': transcription})
    except Exception as e:
        print("Error:", str(e))  # Debugging log
        return jsonify({'error': 'An error occurred while processing the audio file.'}), 500
    finally:
        # Clean up the temporary file
        if os.path.exists(temp_path):
            os.remove(temp_path)

# Download audio transcript endpoint
@app.route('/download-audio-docx', methods=['GET'])
def download_audio_docx():
    docx_path = '/tmp/audio.docx'
    if not os.path.exists(docx_path):
        return jsonify({'error': 'No document found'}), 404
    return send_file(docx_path, as_attachment=True)

@app.route('/tts', methods=['POST'])
def tts():
    data = request.get_json()
    text = data.get('text', '')
    if not text:
        return jsonify({'error': 'No text provided'}), 400

    # Call OpenAI TTS API
    response = openai.audio.speech.create(
        model="tts-1",  # or "tts-1-hd"
        voice="alloy",   # or "nova", "shimmer", "echo", etc.
        input=text
    )
    audio_bytes = response.content

    # Return as a file-like object
    return send_file(
        io.BytesIO(audio_bytes),
        mimetype='audio/mpeg',
        as_attachment=False,
        download_name='speech.mp3'
    )

if __name__ == '__main__':
    app.run(host='127.0.0.1', port=3000, debug=True)